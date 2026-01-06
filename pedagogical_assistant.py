import streamlit as st
import os
import pandas as pd
import io
import re 
import random
import string
import json
from google import genai
# Importamos la clase de Error específica para capturarla
from google.genai.errors import APIError 
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor
# Necesitamos 'parse' para la nueva función de Word
from docx.text.paragraph import Paragraph
from docx.table import _Cell

# --- 👇 NUEVOS IMPORTS NECESARIOS PARA EL COLOR DE FONDO 👇 ---
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
# --------------------------------------------------------------

# =========================================================================
# === I. FUNCIÓN DE PROPUESTAS (Pestaña 1) ===
# =========================================================================

def generate_ai_suggestions(critical_comp_info):
    """
    Genera un plan de acción (propuestas de mejora) usando la IA.
    Usa el modelo 'pro' para alta consistencia de formato.
    """
    
    if client is None:
        return "⚠️ **Error de Configuración de IA:** El cliente de Gemini no se pudo inicializar."
        
    # Extraemos las variables dinámicas
    grado = critical_comp_info['grado']
    nivel = critical_comp_info['nivel']
    area = critical_comp_info['area']
    competencia = critical_comp_info['nombre']
    analisis = critical_comp_info['analisis'] 
    
    prompt = f"""
    Elabora un plan de acción pedagógico completo y práctico para mejorar la competencia "{competencia}" en el área de "{area}" para estudiantes de {grado} de {nivel}.

    Basado en este diagnóstico: {analisis}

    **REGLAS OBLIGATORIAS DE FORMATO (no las ignores, o no se aceptará la respuesta):**
    1. Responde **exclusivamente** con una tabla Markdown. NO pongas texto antes, después, introducciones o conclusiones.
    2. Usa **exactamente** estas 3 columnas:  
       | Acción Concreta | Indicadores de Mejora | Evidencias Esperadas |
       |-----------------|------------------------|----------------------|
    3. Genera **exactamente 5 acciones concretas** (no más, no menos). Cada acción debe ser realista, fácil de implementar por un docente y enfocada en estudiantes con dificultades (niveles B y C).
    4. Usa texto claro y directo dentro de las celdas. Si necesitas destacar algo, usa cursivas (*texto*) o mayúsculas, pero NO uses **negritas** (asteriscos dobles).
    5. Añade un diseño visual:  
       - Encabezados con fondo azul oscuro y texto blanco (usa CSS: <thead style="background-color: #113770; color: white;">).
       - Bordes visibles y gruesos en la tabla.
       - Alterna filas con fondo gris claro (#f8f9fa) para mejor legibilidad (usa <tr style="background-color: #f8f9fa;"> en filas pares).
    6. La tabla debe verse profesional y fácil de copiar.

    Comienza directamente con la tabla <table> en Markdown, sin nada más.
    """
    
    try:
        response = client.models.generate_content(
            model='models/gemini-2.5-flash-preview-09-2025', 
            contents=prompt,
        )
        return response.text
    except APIError as e: 
        return f"❌ **Error al contactar la IA:** Se produjo un error en la API de Google (Código: {e}). Revisa tu clave y la cuota de uso."
    except Exception as e:
        return f"❌ **Error desconocido:** {e}"


# =========================================================================
# === II-A. FUNCIÓN DE EXPORTACIÓN A WORD (Propuestas) ===
# =========================================================================
def generate_docx_report(analisis_results, sheet_name, selected_comp_limpio, ai_report_text):
    document = Document()
    result = analisis_results[sheet_name]
    general_data = result.get('generalidades', {})
    grado = general_data.get('grado', 'Desconocido')
    nivel = general_data.get('nivel', 'Desconocido')
    
    document.add_heading(f'INFORME DE PROPUESTAS PEDAGÓGICAS', 0)
    document.add_heading('Datos de Contexto y Diagnóstico', level=1)
    document.add_paragraph(f"Nivel/Grado: {nivel} / {grado}")
    document.add_paragraph(f"Área Analizada: {sheet_name}")
    
    p_comp = document.add_paragraph()
    p_comp.add_run(f"Competencia a Abordar: ").bold = True
    p_comp.add_run(selected_comp_limpio)
    
    document.add_heading('Propuestas de Intervención (Generadas por IA)', level=1)
    
    def process_markdown_to_runs(paragraph, text):
        parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                paragraph.add_run(part[2:-2]).bold = True
            elif part.startswith('*') and part.endswith('*'):
                paragraph.add_run(part[1:-1]).italic = True
            else:
                paragraph.add_run(part)

    lines = ai_report_text.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
        if "propuestas de intervención didáctica" in line.lower() or "propuestas de intervención (generadas por ia)" in line.lower():
            continue
        if line.startswith('###'):
            document.add_heading(re.sub(r'^###\s*', '', line).strip(), level=2)
        elif line.startswith('##'):
            document.add_heading(re.sub(r'^##\s*', '', line).strip(), level=1)
        elif line.startswith('#'):
            document.add_heading(re.sub(r'^#s*', '', line).strip(), level=1)
        elif re.match(r'^\d+\.', line):
            paragraph = document.add_paragraph(style='List Number')
            cleaned_line = re.sub(r'^\d+\.\s*', '', line).strip()
            process_markdown_to_runs(paragraph, cleaned_line)
        elif line.startswith('*'):
            paragraph = document.add_paragraph(style='List Bullet')
            cleaned_line = re.sub(r'^\*\s*', '', line).strip()
            process_markdown_to_runs(paragraph, cleaned_line)
        else:
            if line: 
                paragraph = document.add_paragraph()
                process_markdown_to_runs(paragraph, line)
                
    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer



# =========================================================================
# === III. FUNCIÓN PRINCIPAL LLAMADA DESDE APP.PY (Propuestas) ===
# =========================================================================
def generate_suggestions(analisis_results, selected_sheet_name, selected_comp_limpio):
    sheet_name = selected_sheet_name
    result = analisis_results[sheet_name]
    
    st.markdown("### 📋 Informe de Intervención Pedagógica")

    if 'error' in result:
        st.error(f"No se puede generar el informe debido al error en el análisis de la hoja '{sheet_name}'.")
        return "Error en el análisis."

    target_comp_data = None
    for original_name, comp_data in result['competencias'].items():
        if comp_data['nombre_limpio'] == selected_comp_limpio:
            target_comp_data = comp_data
            break
            
    if not target_comp_data:
        st.error(f"Error: No se encontró la competencia '{selected_comp_limpio}' en los datos de análisis.")
        return "Error: Competencia no encontrada."

    counts = target_comp_data['conteo_niveles']
    total = target_comp_data['total_evaluados']
    
    if total == 0:
        st.info(f"La competencia '{selected_comp_limpio}' no tiene estudiantes evaluados.")
        return "Error: Sin evaluados."
        
    c_count = counts.get('C', 0)
    b_count = counts.get('B', 0)
    c_percentage = (c_count / total) * 100
    
    critical_comp_info = {
        'area': sheet_name,
        'nombre': selected_comp_limpio,
        'analisis': f"El {c_percentage:.1f}% de los estudiantes ({c_count} estudiantes) se encuentra en el Nivel C. El {((b_count+c_count)/total)*100:.1f}% ({b_count+c_count} estudiantes) está en nivel de inicio (B o C).",
        'grado': result['generalidades'].get('grado'),
        'nivel': result['generalidades'].get('nivel')
    }

    st.markdown(f"**Diagnóstico:** El área de **{critical_comp_info['area']}** en **{critical_comp_info['grado']}** requiere atención especial en la competencia:")
    st.markdown(f"#### ⚠️ {critical_comp_info['nombre']}")
    st.markdown(f"**Análisis de Dificultad:** {critical_comp_info['analisis']}")
    st.markdown("---")

    with st.spinner("🧠 Generando propuestas pedagógicas con Inteligencia Artificial..."):
        ai_response_text = generate_ai_suggestions(critical_comp_info)
        return ai_response_text 



# =========================================================================
# === V. GENERADOR DE INFORME DEL ESTUDIANTE (Word con Colores) ===
# =========================================================================
def generar_reporte_estudiante(nombre_estudiante, total_conteo, desglose_areas):
    """
    Genera un informe individual en Word con formato semáforo (colores).
    """
    document = Document()
    
    # --- ESTILOS ---
    style = document.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    # --- FUNCIÓN INTERNA PARA COLOR (Para pintar celdas en Word) ---
    def set_cell_shading(cell, fill_color):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), fill_color)
        tcPr.append(shd)

    # 1. ENCABEZADO
    # Creamos el título pero accedemos a su "run" (el texto) para cambiarle el tamaño
    h1 = document.add_heading('INFORME DE PROGRESO DEL APRENDIZAJE', 0)
    h1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # AJUSTE DE TAMAÑO (Arial 18)
    run = h1.runs[0]
    run.font.name = 'Arial'
    run.font.size = Pt(18)  # <--- AQUÍ ESTÁ EL EQUIVALENTE A ARIAL 18
    run.font.color.rgb = RGBColor(0, 0, 0) # Aseguramos color negro
    
    document.add_paragraph(f"Estudiante: {nombre_estudiante}")
    document.add_paragraph("Fecha de emisión: _______________________")
    document.add_paragraph("")

    # 2. SEMÁFORO ACADÉMICO (Tabla de Resumen)
    document.add_heading('1. Resumen de Logros (Semáforo)', level=1)
    
    table = document.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    
    # Encabezados
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'NIVEL DE LOGRO'
    hdr_cells[1].text = 'CANTIDAD DE ÁREAS'
    for cell in hdr_cells: 
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, "D9D9D9") # Gris claro para encabezado

    # Datos del semáforo
    data = [
        ("LOGRO DESTACADO (AD)", total_conteo['AD'], "C6EFCE"), # Verde Claro
        ("LOGRO ESPERADO (A)", total_conteo['A'], "E7F3FF"),   # Azul Claro
        ("EN PROCESO (B)", total_conteo['B'], "FFEB9C"),       # Amarillo
        ("EN INICIO (C)", total_conteo['C'], "FFC7CE")         # Rojo Claro
    ]

    for nivel, cantidad, color_hex in data:
        row_cells = table.add_row().cells
        row_cells[0].text = nivel
        row_cells[1].text = str(cantidad)
        
        # Pintamos la celda del nivel
        set_cell_shading(row_cells[0], color_hex)
        row_cells[0].paragraphs[0].runs[0].bold = True
        
        # Centramos la cantidad
        row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_paragraph("")

    # 3. DETALLE DE ÁREAS CRÍTICAS
    if total_conteo['B'] > 0 or total_conteo['C'] > 0:
        document.add_heading('2. Áreas que requieren atención', level=1)
        
        if total_conteo['C'] > 0:
            p = document.add_paragraph()
            run = p.add_run("🛑 EN INICIO (C) - Requiere Recuperación:")
            run.bold = True
            run.font.color.rgb = RGBColor(200, 0, 0) # Rojo oscuro
            
            for area_txt in desglose_areas['C']:
                document.add_paragraph(f"   • {area_txt}", style='List Bullet')
        
        if total_conteo['B'] > 0:
            p = document.add_paragraph()
            run = p.add_run("⚠️ EN PROCESO (B) - Requiere Refuerzo:")
            run.bold = True
            run.font.color.rgb = RGBColor(200, 150, 0) # Naranja oscuro
            
            for area_txt in desglose_areas['B']:
                document.add_paragraph(f"   • {area_txt}", style='List Bullet')

    document.add_paragraph("")

    # 4. RECOMENDACIONES PEDAGÓGICAS (Automáticas)
    document.add_heading('3. Recomendaciones y Compromisos', level=1)
    
    recomendacion = ""
    if total_conteo['C'] > 0:
        recomendacion = "El estudiante requiere un mayor acompañamiento para consolidar los aprendizajes en las áreas señaladas. Se sugiere reforzar los hábitos de estudio en casa y mantener comunicación constante con los docentes para asegurar su proceso de aprendizaje."
    elif total_conteo['B'] > 0:
        recomendacion = "Va por buen camino. Sugerimos motivar al estudiante a participar más activamente y revisar juntos sus avances semanales para que logre alcanzar el nivel de logro esperado en el corto plazo."
    else:
        recomendacion = "¡Felicitaciones! El estudiante demuestra un alto nivel de compromiso y logro de competencias. Se sugiere mantener la motivación, leer libros de interés y explorar nuevos retos académicos."
    
    document.add_paragraph(recomendacion)
    document.add_paragraph("")
    document.add_paragraph("")

    # 5. FIRMAS
    table_firmas = document.add_table(rows=1, cols=2)
    f_cells = table_firmas.rows[0].cells
    
    p1 = f_cells[0].paragraphs[0]
    p1.add_run("_________________________").bold = True
    p1.add_run("\nAPODERADO(A)")
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p2 = f_cells[1].paragraphs[0]
    p2.add_run("_________________________").bold = True
    p2.add_run("\nDOCENTE / TUTOR")
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Guardar en memoria
    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer

# =========================================================================
# === VII. GENERADOR DE TRIVIA (GAMIFICACIÓN) - VERSIÓN DINÁMICA ===
# =========================================================================

def generar_trivia_juego(tema, grado, area, cantidad):
    """
    Genera preguntas de selección múltiple en formato JSON.
    Cantidad ajustable por el usuario (1-10).
    """
    if client is None:
        return None

    prompt = f"""
    Actúa como un diseñador de videojuegos educativos.
    Crea un set de **{cantidad} PREGUNTAS DE TRIVIA** divertidas y desafiantes sobre el tema: "{tema}" para estudiantes de {grado} ({area}).

    REGLAS DE FORMATO (JSON ESTRICTO):
    Devuelve SOLO un array JSON (lista de objetos) con esta estructura exacta:
    [
        {{
            "pregunta": "¿Texto de la pregunta?",
            "opciones": ["Opción A", "Opción B", "Opción C", "Opción D"],
            "respuesta_correcta": "Opción A",
            "explicacion": "Breve explicación de por qué es la correcta."
        }},
        ... (repetir {cantidad} veces)
    ]

    REGLAS DE CONTENIDO:
    1. Las opciones deben ser plausibles.
    2. La "respuesta_correcta" debe coincidir EXACTAMENTE con una de las "opciones".
    3. Lenguaje adecuado para {grado}.
    """

    try:
        response = client.models.generate_content(
            model='models/gemini-2.5-flash',
            contents=prompt,
            config={'response_mime_type': 'application/json'}
        )
        return response.text
    except Exception as e:
        return None


# =========================================================================
# === VIII. MOTOR DE PUPILETRAS (LOGICA MIXTA: IA + ALGORITMO) ===
# =========================================================================

def generar_palabras_pupiletras(tema, grado, cantidad):
    """
    Paso 1: La IA genera la lista de palabras limpia.
    """
    if client is None:
        return None

    prompt = f"""
    Actúa como experto en didáctica. Genera una lista de {cantidad} palabras clave (sustantivos o verbos) sobre el tema: "{tema}" para estudiantes de {grado}.
    
    REGLAS OBLIGATORIAS:
    1. Las palabras deben estar en MAYÚSCULAS.
    2. SIN TILDES (convierte Á->A, É->E, etc).
    3. SIN ESPACIOS (ej: "SISTEMASOLAR" en vez de "SISTEMA SOLAR").
    4. SIN Ñ (cámbiala por N).
    5. Longitud máxima por palabra: 12 letras.
    
    FORMATO JSON:
    Devuelve SOLO una lista simple de strings:
    ["PALABRAUNO", "PALABRADOS", ...]
    """

    try:
        response = client.models.generate_content(
            model='models/gemini-2.5-flash',
            contents=prompt,
            config={'response_mime_type': 'application/json'}
        )
        return json.loads(response.text)
    except Exception as e:
        return []

def crear_grid_pupiletras(palabras, filas=12, columnas=12):
    """
    Paso 2: Algoritmo Python para colocar las palabras en una matriz 12x12.
    Retorna: (grid, palabras_colocadas)
    """
    # 1. Crear grilla vacía
    grid = [[' ' for _ in range(columnas)] for _ in range(filas)]
    palabras_colocadas = []
    
    # Direcciones: (delta_fila, delta_columna)
    # Horizontal, Vertical, Diagonal, Invertidas
    direcciones = [
        (0, 1), (1, 0), (1, 1), (1, -1), # Normales
        (0, -1), (-1, 0), (-1, -1), (-1, 1) # Invertidas (Mayor dificultad)
    ]

    # Ordenamos palabras de mayor a menor longitud (facilita el encaje)
    palabras.sort(key=len, reverse=True)

    for palabra in palabras:
        colocada = False
        intentos = 0
        
        # Intentamos colocar la palabra 100 veces en posiciones al azar
        while not colocada and intentos < 100:
            intentos += 1
            direccion = random.choice(direcciones)
            fila_inicio = random.randint(0, filas - 1)
            col_inicio = random.randint(0, columnas - 1)
            
            # Chequeamos si cabe
            fila, col = fila_inicio, col_inicio
            cabe = True
            
            for letra in palabra:
                # Verificar limites
                if not (0 <= fila < filas and 0 <= col < columnas):
                    cabe = False
                    break
                # Verificar colisión (casilla vacía o misma letra)
                if grid[fila][col] != ' ' and grid[fila][col] != letra:
                    cabe = False
                    break
                
                fila += direccion[0]
                col += direccion[1]
            
            # Si cabe, la escribimos
            if cabe:
                fila, col = fila_inicio, col_inicio
                coords = [] # Guardamos coordenadas para el frontend interactivo
                for letra in palabra:
                    grid[fila][col] = letra
                    coords.append((fila, col))
                    fila += direccion[0]
                    col += direccion[1]
                
                palabras_colocadas.append({
                    "palabra": palabra,
                    "coords": coords
                })
                colocada = True

    # 3. Rellenar espacios vacíos con letras aleatorias
    letras = string.ascii_uppercase
    grid_completo = [] # Copia para mostrar
    for f in range(filas):
        fila_letras = []
        for c in range(columnas):
            if grid[f][c] == ' ':
                fila_letras.append(random.choice(letras))
            else:
                fila_letras.append(grid[f][c])
        grid_completo.append(fila_letras)

    return grid_completo, palabras_colocadas

# =========================================================================
# === IX. GENERADOR DE DOCX PUPILETRAS (FICHA IMPRIMIBLE) ===
# =========================================================================

def generar_docx_pupiletras(grid, palabras_data, tema, grado):
    """
    Crea un archivo Word con la sopa de letras formateada para imprimir.
    """
    doc = Document()
    
    # 1. TÍTULO Y ENCABEZADO
    titulo = doc.add_heading(f"SOPA DE LETRAS: {tema.upper()}", 0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"Nivel: {grado} | Generado por AulaMetrics").bold = True

    doc.add_paragraph("") # Espacio

    # 2. DIBUJAR LA GRILLA (TABLA)
    filas = len(grid)
    columnas = len(grid[0])
    
    # Creamos tabla
    table = doc.add_table(rows=filas, cols=columnas)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.autofit = False 
    
    # Configuración de celdas
    for i in range(filas):
        for j in range(columnas):
            cell = table.cell(i, j)
            cell.text = grid[i][j]
            
            # Formato de texto (Centrado y Grande)
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.runs[0]
            run.font.size = Pt(14)
            run.font.bold = True
            
            # Ajuste de ancho/alto (para que sea cuadrada)
            cell.width = Inches(0.4)
            cell.height = Inches(0.4)

    doc.add_paragraph("") # Espacio grande
    doc.add_paragraph("")

    # 3. LISTA DE PALABRAS A BUSCAR
    doc.add_heading("Palabras a encontrar:", level=2)
    
    # Usamos una tabla invisible para listar las palabras ordenadamente (3 columnas)
    lista_palabras = [item['palabra'] for item in palabras_data]
    lista_palabras.sort()
    
    num_cols_lista = 3
    num_rows_lista = (len(lista_palabras) + num_cols_lista - 1) // num_cols_lista
    
    list_table = doc.add_table(rows=num_rows_lista, cols=num_cols_lista)
    list_table.style = 'Table Grid' 
    
    idx = 0
    for r in range(num_rows_lista):
        for c in range(num_cols_lista):
            if idx < len(lista_palabras):
                cell = list_table.cell(r, c)
                cell.text = f"⬜ {lista_palabras[idx]}"
                idx += 1

    # 4. GUARDAR EN MEMORIA
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io


# =========================================================================
# === X. MOTOR JUEGO ROBOT (AHORCADO EDUCATIVO V2.0) ===
# =========================================================================

def generar_reto_ahorcado(tema, grado, cantidad):
    """
    Genera una lista de palabras y pistas para el juego del Robot.
    Retorna: list [{'palabra':Str, 'pista':Str}, ...]
    """
    if client is None:
        return []

    prompt = f"""
    Actúa como un diseñador de juegos educativos. 
    Necesito {cantidad} retos distintos para un juego tipo "Ahorcado" sobre el tema: "{tema}" para el grado: "{grado}".
    
    INSTRUCCIONES:
    1. Elige palabras clave (conceptos importantes) relacionadas con el tema.
    2. Las palabras deben ser en MAYÚSCULAS, SIN TILDES y SIN ESPACIOS (Ej: "ECOSISTEMA", no "Ecosistema" ni "Árbol").
    3. Escribe una pista pedagógica clara para cada palabra, adecuada al nivel del estudiante.
    
    FORMATO JSON OBLIGATORIO (Array de objetos):
    [
        {{
            "palabra": "PALABRAUNO",
            "pista": "Texto de la pista uno..."
        }},
        {{
            "palabra": "PALABRADOS",
            "pista": "Texto de la pista dos..."
        }}
    ]
    """

    try:
        response = client.models.generate_content(
            model='models/gemini-2.5-flash',
            contents=prompt,
            config={'response_mime_type': 'application/json'}
        )
        return json.loads(response.text)
    except Exception as e:
        print(f"Error generando ahorcado: {e}")
        return []
