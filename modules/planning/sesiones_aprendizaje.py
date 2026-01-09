import streamlit as st
import os
import pandas as pd
import io
import re 
import random
import string
import json
from google import genai
from google.genai.errors import APIError 
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor
from docx.text.paragraph import Paragraph
from docx.table import _Cell
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# =========================================================================
# === 1. FUNCIONES AUXILIARES ===
# =========================================================================
ISOTIPO_PATH = "assets/isotipo.png"
RUTA_ESTANDARES = "assets/Estandares de aprendizaje.xlsx" 

@st.cache_data(ttl=3600)
def cargar_datos_pedagogicos():
    try:
        df_generalidades = pd.read_excel(RUTA_ESTANDARES, sheet_name="Generalidades")
        df_ciclos = pd.read_excel(RUTA_ESTANDARES, sheet_name="Cicloseducativos")
        df_desc_sec = pd.read_excel(RUTA_ESTANDARES, sheet_name="Descriptorsecundaria")
        df_desc_prim = pd.read_excel(RUTA_ESTANDARES, sheet_name="Descriptorprimaria")
        
        df_generalidades['NIVEL'] = df_generalidades['NIVEL'].ffill()
        df_ciclos['ciclo'] = df_ciclos['ciclo'].ffill()
        
        columna_estandar = "DESCRIPCI√ìN DE LOS NIVELES DEL DESARROLLO DE LA COMPETENCIA"
        
        cols_to_fill_prim = ['√Årea', 'Competencia', 'Ciclo', columna_estandar]
        cols_to_fill_sec = ['√Årea', 'Competencia', 'Ciclo', columna_estandar]
        
        df_desc_prim[cols_to_fill_prim] = df_desc_prim[cols_to_fill_prim].ffill()
        df_desc_sec[cols_to_fill_sec] = df_desc_sec[cols_to_fill_sec].ffill()
        
        return df_generalidades, df_ciclos, df_desc_sec, df_desc_prim
    
    except FileNotFoundError:
        st.error(f"Error: No se encontr√≥ el archivo en la ruta: {RUTA_ESTANDARES}")
        return None, None, None, None
    except Exception as e:
        st.error(f"Ocurri√≥ un error al leer el archivo Excel: {e}")
        return None, None, None, None

# =========================================================================
# === 2. CONFIGURACI√ìN GLOBAL DE LA IA ===
# =========================================================================

try:
    gemini_key = st.secrets['gemini']['api_key']
    client = genai.Client(api_key=gemini_key)
except KeyError:
    st.error("Error de Configuraci√≥n: No se encontr√≥ la clave de Gemini en st.secrets.")
    client = None
except Exception as e:
    st.error(f"Error al inicializar el cliente de Gemini: {e}")
    client = None

# =========================================================================
# === 3. FUNCI√ìN DE GENERACI√ìN DE SESI√ìN (Pesta√±a 3) ===
# =========================================================================
def generar_sesion_aprendizaje(nivel, grado, ciclo, area, competencias_lista, capacidades_lista, estandar_texto, tematica, tiempo, region=None, provincia=None, distrito=None, instrucciones_docente=None):
    if client is None:
        return "Error: No se pudo inicializar el cliente de Gemini."

    # Convertir listas a texto formateado
    competencias_str = "\n".join(f"- {comp}" for comp in competencias_lista)
    capacidades_str = "\n".join(f"- {cap}" for cap in capacidades_lista)

    # CONTEXTO GEOGR√ÅFICO
    contexto_str = ""
    if region and region.strip():
        contexto_str = f"""
- **Regi√≥n:** {region}
- **Provincia:** {provincia}
- **Distrito:** {distrito}
DEBES usar estos datos para generar ejemplos relevantes en las tablas.
"""
    else:
        contexto_str = "No hay contexto geogr√°fico. Usa ejemplos generales."

    # INSTRUCCIONES ADICIONALES
    instrucciones_str = ""
    if instrucciones_docente and instrucciones_docente.strip():
        instrucciones_str = f"""
**INSTRUCCIONES ADICIONALES DEL DOCENTE:** {instrucciones_docente}
Modifica todas las tablas para cumplir esto estrictamente.
"""
    else:
        instrucciones_str = "No hay instrucciones adicionales. Genera una sesi√≥n est√°ndar."

    # MEN√ö DE METODOLOG√çAS
    menu_metodologias = """
1. Aprendizaje Basado en Problemas (ABP)
2. Aprendizaje Basado en Indagaci√≥n (Indagaci√≥n Cient√≠fica)
3. Aprendizaje Colaborativo / Cooperativo
4. Gamificaci√≥n (Uso de mec√°nicas de juego)
5. Estudio de Casos
6. Aula Invertida (Flipped Classroom)
"""

    # Prompt estricto con salida JSON
    prompt = f"""
Act√∫a como docente experto en dise√±o curricular peruano.

**ESTRATEGIA PEDAG√ìGICA:**
Analiza Grado ({grado}), √Årea ({area}), Tema ({tematica}). Elige UNA metodolog√≠a de esta lista:
{menu_metodologias}

**MANDATO DE ALTA DEMANDA COGNITIVA:**
En 'desarrollo', incluye al menos una actividad de razonamiento complejo, creatividad o pensamiento cr√≠tico. Enf√≥cate en lo que el estudiante HACE.

**DATOS:**
- Nivel: {nivel}
- Grado: {grado}
- Ciclo: {ciclo}
- √Årea: {area}
- Tema: "{tematica}"
- Duraci√≥n: {tiempo} minutos
- Competencia(s): {competencias_str}
- Capacidad(es): {capacidades_str}
- Est√°ndar(es): "{estandar_texto}"
- Contexto geogr√°fico: {contexto_str}
{instrucciones_str}

**REGLAS OBLIGATORIAS - NO ROMPERLAS:**
1. Responde **SOLO** con JSON v√°lido.
2. NO texto libre, NO introducci√≥n, NO conclusi√≥n.
3. Estructura JSON exacta:
{{
  "titulo": "T√≠tulo creativo de la sesi√≥n",
  "objetivos": ["Objetivo 1", "Objetivo 2", ...],
  "metodologia_elegida": "Nombre de la metodolog√≠a elegida",
  "inicio": [
    {{"actividad": "...", "tiempo": "...", "recursos": "...", "responsable": "..."}},
    ...
  ],
  "desarrollo": [
    {{"paso": "...", "descripcion": "...", "objetivo": "...", "materiales": "..."}},
    ...
  ],
  "cierre": [
    {{"actividad": "...", "tiempo": "...", "evaluacion": "..."}},
    ...
  ],
  "recursos_generales": [
    {{"recurso": "...", "cantidad": "...", "uso": "..."}},
    ...
  ],
  "evaluacion": [
    {{"criterio": "...", "indicador": "...", "instrumento": "..."}},
    ...
  ]
}}

**Output SOLO JSON v√°lido. Nada m√°s.**
"""

    try:
        response = client.models.generate_content(
            model='gemini-1.5-flash-latest',
            contents=prompt,
            response_mime_type='application/json'   # ‚Üê as√≠, sin "generation_config"
        )
        return response.text.strip()
    except APIError as e:
        return f"Error de API: {e}"
    except Exception as e:
        return f"Error inesperado: {e}"

# =========================================================================
# === 4. EXPORTACI√ìN A WORD INTELIGENTE (Sesi√≥n) - v4.0 COLOR ===
# =========================================================================
def generar_docx_sesion(sesion_markdown_text, area_docente):
    """
    Convierte la sesi√≥n a Word con:
    - Filtro de inicio estricto.
    - Tabla de competencias con ENCABEZADOS DE COLOR.
    - Limpieza y formato de texto reparados.
    """
    document = Document()
    
    # --- ESTILOS B√ÅSICOS ---
    style = document.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    # --- FUNCI√ìN TRUCO PARA PINTAR CELDAS (XML) ---
    def set_cell_shading(cell, fill_color):
        """
        Pinta el fondo de una celda. 
        fill_color: C√≥digo Hexadecimal sin # (ej: 'E7F3FF').
        """
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), fill_color)
        tcPr.append(shd)

    # --- HELPERS DE TEXTO ---
    def clean_markdown_symbols(text):
        return text.replace('>', '').strip()

    def clean_asterisks(text):
        return text.replace('**', '').replace('*', '').strip()

    def process_formatted_text(paragraph, text):
        text = clean_markdown_symbols(text)
        parts = re.split(r'(\*\*.*?\*\*)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                clean_part = part[2:-2]
                run = paragraph.add_run(clean_part)
                run.bold = True
            else:
                paragraph.add_run(part)

    def add_bullet(paragraph, text, style='List Bullet'):
        clean_text = re.sub(r'^[\*\-\+]\s*', '', text)
        paragraph.style = style
        process_formatted_text(paragraph, clean_text)

    lines = sesion_markdown_text.split('\n')
    
    # --- VARIABLES DE ESTADO ---
    printing_started = False 
    in_competencies_section = False
    table = None
    curr_comp = ""
    curr_caps = []
    curr_crits = []
    capture_mode = 0 

    def flush_row(tbl, c, caps, crits):
        if not c and not caps and not crits: return
        row = tbl.add_row()
        row.cells[0].text = clean_asterisks(c)
        if caps:
            row.cells[1].paragraphs[0].text = "" 
            for cap in caps:
                p = row.cells[1].add_paragraph()
                add_bullet(p, cap)
        if crits:
            row.cells[2].paragraphs[0].text = ""
            for crit in crits:
                p = row.cells[2].add_paragraph()
                add_bullet(p, crit)

    # --- BUCLE PRINCIPAL ---
    for line in lines:
        line = line.strip()
        if not line: continue

        if line.upper().startswith("### SESI√ìN DE APRENDIZAJE") and not printing_started:
            printing_started = True
            p = document.add_heading('SESI√ìN DE APRENDIZAJE', level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue 
        
        if not printing_started: continue 

        if re.match(r'^(\*\*)?(I|II|III|IV|V|VI|VII)\.', line):
            if in_competencies_section and table:
                flush_row(table, curr_comp, curr_caps, curr_crits)
                curr_comp, curr_caps, curr_crits = "", [], []
                in_competencies_section = False

            clean_title = line.replace('**', '').strip()
            document.add_heading(clean_title, level=1)

            if "COMPETENCIAS" in line.upper():
                in_competencies_section = True
                table = document.add_table(rows=1, cols=3)
                table.style = 'Table Grid'
                
                # --- CONFIGURACI√ìN DE ENCABEZADOS CON COLOR ---
                hdr_row = table.rows[0]
                hdr = hdr_row.cells
                
                headers = ['COMPETENCIA', 'CAPACIDAD', 'CRITERIOS DE EVALUACI√ìN']
                # Color Azul Suave (Hex: D9EAD3 es verde suave, usaremos E7F3FF para azul suave)
                color_header = "E7F3FF" 
                
                for i, text in enumerate(headers):
                    hdr[i].text = text
                    # ¬°Aqu√≠ aplicamos el color!
                    set_cell_shading(hdr[i], color_header) 
                    # Ponemos negrita
                    hdr[i].paragraphs[0].runs[0].bold = True
                # ---------------------------------------------
            
            continue

        # --- L√ìGICA DE TABLA (INTACTA) ---
        if in_competencies_section:
            if "---" in line:
                flush_row(table, curr_comp, curr_caps, curr_crits)
                curr_comp, curr_caps, curr_crits = "", [], []
                capture_mode = 0
                continue
            if "Competencia:" in line: 
                if curr_comp: 
                     flush_row(table, curr_comp, curr_caps, curr_crits)
                     curr_caps, curr_crits = [], []
                text_parts = line.split("ompetencia:") 
                if len(text_parts) > 1:
                    curr_comp = text_parts[1].replace('**', '').strip()
                capture_mode = 0
            elif "Capacidades:" in line: capture_mode = 1
            elif "Criterios" in line and "Evaluaci√≥n" in line: capture_mode = 2
            elif line.startswith('-') or line.startswith('*'):
                content = re.sub(r'^[\*\-]\s*', '', line).strip()
                if capture_mode == 1: curr_caps.append(content)
                elif capture_mode == 2: curr_crits.append(content)
            continue 

        # --- CONTENIDO NORMAL ---
        if line.startswith("###"):
            clean_h = line.replace('###', '').strip()
            document.add_heading(clean_h, level=2)
        elif line.startswith('**') and line.endswith('**') and ":" in line:
             p = document.add_paragraph()
             clean_line = line.replace('**', '').strip()
             run = p.add_run(clean_line)
             run.bold = True
        elif line.startswith('* ') or line.startswith('- '):
            p = document.add_paragraph()
            add_bullet(p, line, style='List Bullet')
        elif line.startswith('>'):
            clean_line = line.replace('>', '').strip()
            if clean_line.startswith('*') or clean_line.startswith('-'):
                p = document.add_paragraph()
                add_bullet(p, clean_line)
            else:
                p = document.add_paragraph()
                process_formatted_text(p, clean_line)
        elif re.match(r'^\d+\.', line):
            p = document.add_paragraph(style='List Number')
            clean_text = re.sub(r'^\d+\.\s*', '', line)
            process_formatted_text(p, clean_text)
        elif line.startswith('_'):
            p = document.add_paragraph(line)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            p = document.add_paragraph()
            process_formatted_text(p, line)

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer


def run():
    st.header("üß† Asistente Pedag√≥gico")

    # Funci√≥n de limpieza local
    def limpiar_resultados():
        keys_to_clear = ['sesion_generada', 'docx_bytes', 'doc_buffer']
        for key in keys_to_clear:
            if key in st.session_state:
                del st.session_state[key]

    tipo_herramienta = st.radio(
        "01. Selecciona la herramienta que deseas usar:",
        options=["Sesi√≥n de aprendizaje", "Unidad de aprendizaje", "Planificaci√≥n Anual"],
        index=0,
        horizontal=True,
        key="asistente_tipo_herramienta",
        on_change=limpiar_resultados
    )
    st.markdown("---")

    if tipo_herramienta == "Sesi√≥n de aprendizaje":
        st.subheader("Generador de Sesi√≥n de Aprendizaje")

        # Carga de datos pedag√≥gicos (aseg√∫rate de tener esta funci√≥n o importarla)
        df_gen, df_cic, df_desc_sec, df_desc_prim = cargar_datos_pedagogicos()

        if df_gen is None:
            st.error("Error cr√≠tico: No se pudieron cargar los est√°ndares.")
        else:
            st.subheader("Paso 1: Selecciona el Nivel")
            niveles = df_gen['NIVEL'].dropna().unique()

            nivel_sel = st.selectbox(
                "Nivel",
                options=niveles,
                index=None,
                placeholder="Elige una opci√≥n...",
                key="asistente_nivel_sel",
                label_visibility="collapsed",
                on_change=limpiar_resultados
            )

            st.subheader("Paso 2: Selecciona el Grado")
            grados_options = []
            if st.session_state.asistente_nivel_sel:
                grados_options = df_gen[df_gen['NIVEL'] == st.session_state.asistente_nivel_sel]['GRADO CORRESPONDIENTE'].dropna().unique()

            grado_sel = st.selectbox(
                "Grado",
                options=grados_options,
                index=None,
                placeholder="Elige un Nivel primero...",
                disabled=(not st.session_state.asistente_nivel_sel),
                key="asistente_grado_sel",
                label_visibility="collapsed",
                on_change=limpiar_resultados
            )

            st.subheader("Paso 3: Selecciona el √Årea")
            areas_options = []
            df_hoja_descriptor = None
            if st.session_state.asistente_grado_sel:
                if st.session_state.asistente_nivel_sel == "SECUNDARIA":
                    df_hoja_descriptor = df_desc_sec
                elif st.session_state.asistente_nivel_sel == "PRIMARIA":
                    df_hoja_descriptor = df_desc_prim
                if df_hoja_descriptor is not None:
                    areas_options = df_hoja_descriptor['√Årea'].dropna().unique()
            area_sel = st.selectbox("√Årea", options=areas_options, index=None, placeholder="Elige un Grado primero...", disabled=(not st.session_state.asistente_grado_sel), key="asistente_area_sel", label_visibility="collapsed")

            st.subheader("Paso 4: Selecciona la(s) Competencia(s)")
            competencias_options = []
            if st.session_state.asistente_area_sel and (df_hoja_descriptor is not None):
                competencias_options = df_hoja_descriptor[
                    df_hoja_descriptor['√Årea'] == st.session_state.asistente_area_sel
                ]['Competencia'].dropna().unique()
            competencias_sel = st.multiselect("Competencia(s)", options=competencias_options, placeholder="Elige un √Årea primero...", disabled=(not st.session_state.asistente_area_sel), key="asistente_competencias_sel", label_visibility="collapsed")

            form_disabled = not st.session_state.asistente_competencias_sel
            st.markdown("---")
            st.subheader("Paso 5: Contextualizaci√≥n (Opcional)")
            contexto_toggle = st.toggle("¬øDesea contextualizar su sesi√≥n?", key="asistente_contexto", disabled=form_disabled)
            with st.form(key="session_form"):
                if st.session_state.asistente_contexto:
                    st.info("La IA usar√° estos datos para generar ejemplos y situaciones relevantes.")
                    region_sel = st.text_input("Regi√≥n de su I.E.", placeholder="Ej: Lambayeque", disabled=form_disabled)
                    provincia_sel = st.text_input("Provincia de su I.E.", placeholder="Ej: Chiclayo", disabled=form_disabled)
                    distrito_sel = st.text_input("Distrito de su I.E.", placeholder="Ej: Monsef√∫", disabled=form_disabled)
                else:
                    region_sel = None
                    provincia_sel = None
                    distrito_sel = None

                st.markdown("---")
                st.subheader("Paso 6: Instrucciones Adicionales (Opcional)")
                instrucciones_sel = st.text_area("Indica un enfoque espec√≠fico para la IA", placeholder="Ej: Quiero reforzar el c√°lculo de porcentajes...", max_chars=500, disabled=form_disabled)
                st.markdown("---")
                st.subheader("Paso 7: Detalles de la Sesi√≥n")
                tema_sel = st.text_input("Escribe el tema o tem√°tica a tratar", placeholder="Ej: El sistema solar...", disabled=form_disabled)
                tiempo_sel = st.selectbox("Selecciona la duraci√≥n de la sesi√≥n", options=["90 minutos", "180 minutos"], index=None, placeholder="Elige una opci√≥n...", disabled=form_disabled)

                submitted = st.form_submit_button("Generar Sesi√≥n de Aprendizaje", disabled=form_disabled)

                if submitted:
                    if not tema_sel or not tiempo_sel:
                        st.error("Por favor, completa los campos del Paso 7.")
                        st.session_state.sesion_generada = None
                        st.session_state.docx_bytes = None
                    else:
                        with st.spinner("ü§ñ Generando tu sesi√≥n de aprendizaje..."):
                            try:
                                nivel = st.session_state.asistente_nivel_sel
                                grado = st.session_state.asistente_grado_sel
                                area = st.session_state.asistente_area_sel
                                competencias = st.session_state.asistente_competencias_sel
                                tema = tema_sel
                                tiempo = tiempo_sel

                                ciclo_float = df_cic[df_cic['grados que corresponde'] == grado]['ciclo'].iloc[0]
                                ciclo_encontrado = int(ciclo_float)
                                datos_filtrados = df_hoja_descriptor[(df_hoja_descriptor['√Årea'] == area) & (df_hoja_descriptor['Competencia'].isin(competencias))]
                                capacidades_lista = datos_filtrados['capacidad'].dropna().unique().tolist()
                                columna_estandar_correcta = "DESCRIPCI√ìN DE LOS NIVELES DEL DESARROLLO DE LA COMPETENCIA"
                                estandares_lista = datos_filtrados[columna_estandar_correcta].dropna().unique().tolist()
                                estandar_texto_completo = "\n\n".join(estandares_lista)
                                
                                # Aqu√≠ usamos las funciones del mismo archivo
                                sesion_generada = generar_sesion_aprendizaje(
                                    nivel=nivel, grado=grado, ciclo=str(ciclo_encontrado), area=area,
                                    competencias_lista=competencias, capacidades_lista=capacidades_lista,
                                    estandar_texto=estandar_texto_completo, tematica=tema, tiempo=tiempo,
                                    region=region_sel, provincia=provincia_sel, distrito=distrito_sel,
                                    instrucciones_docente=instrucciones_sel
                                )
                                docx_bytes = generar_docx_sesion(sesion_markdown_text=sesion_generada, area_docente=area)

                                st.session_state.sesion_generada = sesion_generada
                                st.session_state.docx_bytes = docx_bytes
                                st.session_state.tema_sesion = tema_sel
                                st.success("¬°Sesi√≥n generada!")
                            except Exception as e:
                                st.error(f"Error: {e}")
                                st.session_state.sesion_generada = None

        # MOSTRAR RESULTADOS
        if st.session_state.get('sesion_generada'):
            try:
                # Intentamos parsear como JSON
                data = json.loads(st.session_state.sesion_generada)
                
                # T√≠tulo principal
                st.markdown(f"### {data.get('titulo', 'Sesi√≥n generada')}")
                
                # Objetivos (lista)
                st.markdown("#### Objetivos de la Sesi√≥n")
                objetivos = data.get('objetivos', [])
                if objetivos:
                    st.markdown("\n".join(f"- {obj}" for obj in objetivos))
                else:
                    st.markdown("No se definieron objetivos.")
                
                # Metodolog√≠a
                st.markdown(f"#### Metodolog√≠a Elegida\n{data.get('metodologia_elegida', 'No especificada')}")
                
                # Tabla de Inicio
                st.markdown("#### Inicio")
                inicio = data.get('inicio', [])
                if inicio:
                    st.table(pd.DataFrame(inicio))
                else:
                    st.markdown("No hay actividades de inicio definidas.")
                
                # Tabla de Desarrollo
                st.markdown("#### Desarrollo")
                desarrollo = data.get('desarrollo', [])
                if desarrollo:
                    st.table(pd.DataFrame(desarrollo))
                else:
                    st.markdown("No hay actividades de desarrollo definidas.")
                
                # Tabla de Cierre
                st.markdown("#### Cierre")
                cierre = data.get('cierre', [])
                if cierre:
                    st.table(pd.DataFrame(cierre))
                else:
                    st.markdown("No hay actividades de cierre definidas.")
                
                # Recursos Generales
                st.markdown("#### Recursos Generales")
                recursos = data.get('recursos_generales', [])
                if recursos:
                    st.table(pd.DataFrame(recursos))
                else:
                    st.markdown("No hay recursos generales definidos.")
                
                # Evaluaci√≥n
                st.markdown("#### Evaluaci√≥n")
                evaluacion = data.get('evaluacion', [])
                if evaluacion:
                    st.table(pd.DataFrame(evaluacion))
                else:
                    st.markdown("No hay criterios de evaluaci√≥n definidos.")
                
                st.success("¬°Sesi√≥n generada con √©xito!")
                
            except json.JSONDecodeError:
                # Si no es JSON v√°lido (por si algo falla), fallback a Markdown crudo
                st.markdown("#### Resultado (formato crudo)")
                st.markdown(st.session_state.sesion_generada)
                st.warning("La respuesta no fue JSON v√°lido. Se muestra en formato crudo.")
            
            except Exception as e:
                st.error(f"Error al mostrar la sesi√≥n: {e}")
                st.markdown(st.session_state.sesion_generada)

            # ZONA DE DESCARGA (queda igual)
            if st.session_state.get('docx_bytes'):
                st.download_button(
                    label="üìÑ Descargar Sesi√≥n en Word",
                    data=st.session_state.docx_bytes,
                    file_name="Sesion_Aprendizaje.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
            else:
                st.warning("‚ö†Ô∏è El archivo expir√≥. Por favor genera la sesi√≥n de nuevo.")

    elif tipo_herramienta == "Unidad de aprendizaje":
        st.info("Funci√≥n de Unidades de Aprendizaje (Pr√≥ximamente).")

    elif tipo_herramienta == "Planificaci√≥n Anual":
        st.info("Funci√≥n de Planificaci√≥n Anual (Pr√≥ximamente).")
